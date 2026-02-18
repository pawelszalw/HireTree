import io
import re
import hashlib
from fastapi import UploadFile

ALLOWED_EXTENSIONS = {".pdf", ".docx"}


async def extract_text(file: UploadFile) -> str:
    filename = (file.filename or "").lower()

    if not any(filename.endswith(ext) for ext in ALLOWED_EXTENSIONS):
        raise ValueError("Unsupported file type. Only PDF and DOCX are accepted.")

    raw_bytes = await file.read()

    if filename.endswith(".pdf"):
        return _extract_pdf(raw_bytes)

    return _extract_docx(raw_bytes)


def _extract_pdf(raw_bytes: bytes) -> str:
    import fitz  # pymupdf

    doc = fitz.open(stream=raw_bytes, filetype="pdf")
    pages = [page.get_text() for page in doc]
    return "\n".join(pages)


def _extract_docx(raw_bytes: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(raw_bytes))

    parts = [para.text for para in document.paragraphs]

    # Include text from table cells — python-docx doesn't include these in paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


# Regex patterns for PII removal
_RE_EMAIL = re.compile(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+")

_RE_PHONE = re.compile(
    r"(\+?\d{1,3}[\s.\-]?)?(\(?\d{2,3}\)?[\s.\-]?)(\d{2,4}[\s.\-]?){2,4}\d{2,4}"
)

# Matches two or three capitalised words at the start of a line (likely full name)
_RE_NAME = re.compile(
    r"^([A-Z\u00C0-\u024F][a-zA-Z\u00C0-\u024F'\-]+ ){1,2}"
    r"[A-Z\u00C0-\u024F][a-zA-Z\u00C0-\u024F'\-]+",
    re.MULTILINE,
)


def anonymize(text: str) -> str:
    text = _RE_NAME.sub("[NAME REDACTED]", text, count=1)
    text = _RE_EMAIL.sub("[EMAIL REDACTED]", text)
    text = _RE_PHONE.sub("[PHONE REDACTED]", text)
    return text


def fingerprint(text: str) -> str:
    return hashlib.md5(text.encode("utf-8")).hexdigest()
